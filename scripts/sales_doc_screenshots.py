# -*- coding: utf-8 -*-
"""产品说明书截图加载与 Word 插入辅助。"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SCREENSHOT_DIR = ROOT / "docs" / "sales" / "screenshots"
MANIFEST_PATH = SCREENSHOT_DIR / "manifest.json"

COLOR_MUTED = RGBColor(0x88, 0x88, 0x88)


class ScreenshotLoader:
    """按 manifest 配置，在 Word 对应章节后插入截图。"""

    def __init__(self, screenshot_dir: Path | None = None, manifest_path: Path | None = None):
        self.dir = screenshot_dir or SCREENSHOT_DIR
        self.manifest_path = manifest_path or MANIFEST_PATH
        self.sections: dict = {}
        self.inserted: list[str] = []
        self.missing: list[str] = []
        self._load_manifest()

    def _load_manifest(self) -> None:
        if not self.manifest_path.exists():
            return
        data = json.loads(self.manifest_path.read_text(encoding="utf-8"))
        self.sections = data.get("sections", {})

    def insert(self, doc: Document, section_key: str, width_cm: float = 14.0) -> bool:
        """在 doc 当前位置插入截图（若文件存在）。返回是否成功插入。"""
        cfg = self.sections.get(section_key)
        if not cfg:
            return False

        img_path = self.dir / cfg["file"]
        if not img_path.exists():
            self.missing.append(f"{section_key} → {cfg['file']}")
            return False

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Cm(width_cm))

        caption = cfg.get("caption", "")
        if caption:
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cp.add_run(caption)
            cr.font.name = "微软雅黑"
            cr._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            cr.font.size = Pt(9)
            cr.font.color.rgb = COLOR_MUTED
            cp.paragraph_format.space_after = Pt(6)

        doc.add_paragraph()
        self.inserted.append(section_key)
        return True

    def summary(self) -> str:
        lines = [
            f"已插入截图 {len(self.inserted)} 张",
            f"缺失截图 {len(self.missing)} 张",
        ]
        if self.inserted:
            lines.append("  已插入: " + ", ".join(self.inserted))
        if self.missing:
            lines.append("  缺失: " + "; ".join(self.missing))
        return "\n".join(lines)
