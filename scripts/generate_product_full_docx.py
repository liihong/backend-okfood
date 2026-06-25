#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 OK饭 完整产品功能介绍说明书（面向客户）Word 文档。"""

import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts"))
from sales_doc_screenshots import ScreenshotLoader  # noqa: E402

OUT_PATH = ROOT / "docs" / "sales" / "OK饭产品功能介绍说明书.docx"

COLOR_PRIMARY = RGBColor(0x1A, 0x7A, 0x4A)
COLOR_TEXT = RGBColor(0x33, 0x33, 0x33)
COLOR_MUTED = RGBColor(0x66, 0x66, 0x66)


def set_font(run, size: Pt | None = None, bold: bool = False, color: RGBColor | None = None):
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size:
        run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def add_title(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    sizes = {1: Pt(16), 2: Pt(14), 3: Pt(12)}
    colors = {1: COLOR_PRIMARY, 2: COLOR_PRIMARY, 3: COLOR_TEXT}
    r = p.add_run(text)
    set_font(r, size=sizes.get(level, Pt(12)), bold=True, color=colors.get(level, COLOR_TEXT))
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(4)


def add_para(doc: Document, text: str, bold: bool = False, indent: bool = False, center: bool = False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    set_font(r, size=Pt(10.5), bold=bold, color=COLOR_TEXT)
    p.paragraph_format.space_after = Pt(3)


def add_bullet(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = p.add_run(f"{bold_prefix}：")
        set_font(r1, size=Pt(10.5), bold=True, color=COLOR_PRIMARY)
        r2 = p.add_run(text)
        set_font(r2, size=Pt(10.5), color=COLOR_TEXT)
    else:
        r = p.add_run(text)
        set_font(r, size=Pt(10.5), color=COLOR_TEXT)
    p.paragraph_format.space_after = Pt(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        r = hdr_cells[i].paragraphs[0].add_run(h)
        set_font(r, size=Pt(10), bold=True, color=COLOR_PRIMARY)
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            r = cell.paragraphs[0].add_run(cell_text)
            set_font(r, size=Pt(10), color=COLOR_TEXT)
    doc.add_paragraph()


def build_document(shots: ScreenshotLoader | None = None) -> Document:
    doc = Document()
    shots = shots or ScreenshotLoader()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)

    # 封面标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("OK饭（OK Fine）健康餐配送管理系统")
    set_font(r, size=Pt(22), bold=True, color=COLOR_PRIMARY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("产品功能介绍说明书")
    set_font(r, size=Pt(16), bold=True, color=COLOR_TEXT)

    doc.add_paragraph()
    add_para(doc, "文档版本：V1.0", center=True)
    add_para(doc, "适用对象：意向采购客户、业务决策人", center=True)
    add_para(doc, "文档性质：产品功能说明（非技术文档）", center=True)
    doc.add_paragraph()

    # 一、产品概述
    add_title(doc, "一、产品概述", 1)
    add_para(doc, "OK饭（品牌名：OK Fine）是一套专为 健康餐订阅配送 行业打造的一体化经营管理系统。")
    add_para(doc, "系统覆盖从 用户点餐购卡 → 厨房备餐 → 配送履约 → 财务对账 → 营销运营 的全链路，帮助健康餐门店、中央厨房、连锁品牌实现数字化经营，降低人工统计成本，提升配送效率与会员留存。")

    add_title(doc, "您将获得什么？", 2)
    add_table(doc, ["组成部分", "使用方式", "主要使用者"], [
        ["微信小程序", "微信内打开，无需下载 App", "会员用户、门店配送员"],
        ["Web 管理后台", "电脑浏览器登录", "店主、客服、财务人员"],
        ["云端业务系统", "后台自动运行，无需人工干预", "全店业务数据中枢"],
    ])

    add_title(doc, "适用业态", 2)
    add_bullet(doc, "健康餐 / 轻食 / 减脂餐 周订、月订 门店")
    add_bullet(doc, "支持 午餐、晚餐 分餐段运营")
    add_bullet(doc, "同时服务 长期订户 与 偶尔单次点餐 客户")
    add_bullet(doc, "可扩展至 多门店、多品牌 连锁经营")

    # 二、核心价值
    add_title(doc, "二、核心价值", 1)

    add_title(doc, "对门店经营者", 2)
    add_bullet(doc, "一张配送大表 统筹当日全部订单：订户、单次点餐、自提、请假一目了然")
    add_bullet(doc, "财务实收自动汇总，开卡、续费、退款分项清晰，减少手工对账")
    add_bullet(doc, "会员档案集中管理，请假、改地址、充值、退款一站处理")
    add_bullet(doc, "多种配送方式灵活切换：自有骑手、顺丰同城、UU 跑腿")

    add_title(doc, "对会员用户", 2)
    add_bullet(doc, "微信里完成 看菜单、购卡、请假、查订单，操作简单")
    add_bullet(doc, "配送完成后 微信消息提醒，低余额自动提醒续费")
    add_bullet(doc, "支持 抖音团购券兑换 本地会员卡或优惠券")
    add_bullet(doc, "请假、份数、扣次记录 透明可查，增强用户信任")

    add_title(doc, "对配送团队", 2)
    add_bullet(doc, "配送员在小程序内查看 按片区分组的当日任务")
    add_bullet(doc, "一键拨号、地图导航、确认送达")
    add_bullet(doc, "配送费 自动累计，便于月底结算")

    # 三、微信小程序（会员端）
    add_title(doc, "三、微信小程序（会员端）", 1)
    add_para(doc, "会员通过微信小程序即可完成全部日常操作，底部导航为：首页 · 点单 · 订单 · 我的。")

    add_title(doc, "3.1 注册登录", 2)
    add_bullet(doc, "微信 手机号一键登录，无需额外注册")
    add_bullet(doc, "完善昵称、头像、配送偏好等个人资料")

    add_title(doc, "3.2 首页浏览", 2)
    add_bullet(doc, "门店 轮播广告、活动海报")
    add_bullet(doc, "今日推荐菜品 展示")
    add_bullet(doc, "会员卡 / 卡包 购买入口")
    add_bullet(doc, "门店名称、Logo、联系电话等信息")
    shots.insert(doc, "03_home")

    add_title(doc, "3.3 菜单与点餐", 2)
    add_bullet(doc, "查看 本周 / 下周菜单，按日展示每日菜品")
    add_bullet(doc, "支持 午餐 / 晚餐 切换")
    add_bullet(doc, "查看菜品详情：图片、成分、说明等")
    add_bullet(doc, "选择 配送到家 或 门店自提")
    add_bullet(doc, "单次点餐：选菜、选份数、选供餐日、选地址，微信支付下单")
    add_bullet(doc, "会员可用 剩余餐次余额 抵扣符合条件的单次订单")
    shots.insert(doc, "03_menu")

    add_title(doc, "3.4 会员卡与卡包", 2)
    add_bullet(doc, "浏览可购买的 周卡、月卡、次卡 及自定义卡包")
    add_bullet(doc, "在线下单，微信支付 购卡")
    add_bullet(doc, "购卡时可使用 优惠券 抵扣")
    add_bullet(doc, "查看卡包订单详情与支付状态")
    add_bullet(doc, "支持 分享给微信好友")
    shots.insert(doc, "03_card")

    add_title(doc, "3.5 订阅会员专属功能", 2)
    add_table(doc, ["功能", "说明"], [
        ["请假管理", "明天快速请假、指定日期区间请假、取消请假；午餐/晚餐可分别请假"],
        ["份数设置", "设置每日配送份数（如一人两份），次日生效"],
        ["扣次记录", "查看历史配送扣减明细，消费透明"],
        ["恢复配送", "新购卡后设置起送日期与配送地址"],
        ["个性化设置", "暂停/恢复配送、固定周六不配送等"],
    ])
    shots.insert(doc, "03_mine")

    add_title(doc, "3.6 地址管理", 2)
    add_bullet(doc, "新增、编辑、删除 多个收货地址")
    add_bullet(doc, "地图选点录入地址")
    add_bullet(doc, "提交时 自动校验 是否在门店配送范围内")

    add_title(doc, "3.7 订单中心", 2)
    add_para(doc, "统一管理三类订单，可按状态筛选：")
    add_bullet(doc, "单次点餐订单 — 按日下单的一餐")
    add_bullet(doc, "商城零售订单 — 普通商品购买")
    add_bullet(doc, "卡包订单 — 在线购卡记录")
    add_para(doc, "支持：下单支付、取消未支付订单、查看详情。")

    add_title(doc, "3.8 优惠券", 2)
    add_bullet(doc, "查看 可用、已用、过期 优惠券")
    add_bullet(doc, "下单或购卡时 自动匹配 可用券")
    add_bullet(doc, "有购卡优惠时 主动弹窗提醒")

    add_title(doc, "3.9 抖音团购兑换", 2)
    add_bullet(doc, "在抖音购买的团购券，可在小程序 验券兑换")
    add_bullet(doc, "兑换后可获得：周卡/月卡、会员卡包、优惠券等本地权益")

    add_title(doc, "3.10 消息提醒", 2)
    add_bullet(doc, "送达通知：餐品送达后微信消息提醒")
    add_bullet(doc, "低余额提醒：餐次不足时提醒续费")

    # 四、微信小程序（配送员端）
    add_title(doc, "四、微信小程序（配送员端）", 1)
    add_para(doc, "门店配送员在同一小程序内切换 骑手模式，独立工作台，无需额外安装 App。")

    add_title(doc, "4.1 登录", 2)
    add_bullet(doc, "工号 + PIN 码，或手机号登录")

    add_title(doc, "4.2 今日配送任务", 2)
    add_bullet(doc, "按 配送片区 分组展示当日待送会员")
    add_bullet(doc, "显示：姓名、电话、地址、备注、份数或菜品")
    add_bullet(doc, "区分 待配送 / 已送达 列表")
    add_bullet(doc, "一键拨号、地图导航 至收货地址")
    shots.insert(doc, "04_courier")

    add_title(doc, "4.3 确认送达", 2)
    add_bullet(doc, "订阅会员送达：按每日份数 自动扣减餐次")
    add_bullet(doc, "单次点餐送达：单独确认履约")
    add_bullet(doc, "配送费 自动累计，便于结算")
    add_bullet(doc, "实时统计今日待送 / 已送单量")

    # 五、Web 管理后台
    add_title(doc, "五、Web 管理后台", 1)
    add_para(doc, "管理后台供门店日常经营使用，按角色分配不同权限。")

    add_title(doc, "5.1 角色与权限", 2)
    add_table(doc, ["角色", "适用人员", "主要权限"], [
        ["店主", "门店负责人", "全部功能，含财务、门店配置"],
        ["客服", "日常运营人员", "会员、订单、配送、菜单、营销（不含财务与系统配置）"],
        ["配送专员", "配送管理人员", "配送区域、资质检验、顺丰监控、配送员管理"],
        ["平台管理员", "系统运营方", "多租户、多门店、对接配置"],
    ])

    add_title(doc, "5.2 营业概览", 2)
    add_para(doc, "门店每日经营的 数据驾驶舱：")
    add_bullet(doc, "今日 / 明日备餐 总览：到家配送、自提、请假人数")
    add_bullet(doc, "与 上周同期 对比")
    add_bullet(doc, "片区会员分布 与覆盖情况")
    add_bullet(doc, "当日过期份数、厨房备餐计划")
    add_bullet(doc, "日库存调整：报损、试吃、补偿餐等记录")
    add_bullet(doc, "系统消息提醒：顺丰推单异常、余量不足、开卡待处理等")
    shots.insert(doc, "05_dashboard")

    add_title(doc, "5.3 会员档案库", 2)
    add_bullet(doc, "检索、筛选、导出 会员列表（按片区、套餐、状态等）")
    add_bullet(doc, "查看 / 编辑会员资料：姓名、电话、套餐、余额、起送日等")
    add_bullet(doc, "代客请假、代客改地址")
    add_bullet(doc, "手动充值 / 调整餐次")
    add_bullet(doc, "补偿餐次（客诉处理）")
    add_bullet(doc, "会员退款（退卡预览与执行）")
    add_bullet(doc, "查看操作日志、历史送达记录")
    shots.insert(doc, "05_members")

    add_title(doc, "5.4 开卡工单", 2)
    add_para(doc, "面向 线下收款、微信转账、抖音 等渠道的统一开卡流程：")
    add_bullet(doc, "新建工单：新会员建档开卡 或 老会员续卡")
    add_bullet(doc, "选择卡型、缴费渠道、金额")
    add_bullet(doc, "跟踪状态：未缴 / 已缴 / 已入账 / 已退款 / 已取消")
    add_bullet(doc, "同步微信支付、入账到会员账户（增加餐次并激活计划）")
    add_bullet(doc, "设置 起送日期")
    add_para(doc, "线下收款与小程序在线购卡，最终 统一入账到同一套会员体系，避免两套账。")
    shots.insert(doc, "05_card_orders")

    add_title(doc, "5.5 订单管理", 2)
    add_table(doc, ["订单类型", "业务说明", "主要操作"], [
        ["单次点餐", "按供餐日的散客/会员单餐", "派单、标记送达、改单、取消、退款"],
        ["商城订单", "零售商品订单", "派单、批量派单、标记送达、取消、退款"],
        ["卡包订单", "小程序在线购卡", "查看、退款（实体卡包可派单配送）"],
    ])
    add_para(doc, "通用能力：按日期筛选、批量操作、指派配送员、同步支付状态。")
    shots.insert(doc, "05_orders")

    add_title(doc, "5.6 配送管理", 2)
    add_para(doc, "智能配送大表", bold=True)
    add_bullet(doc, "按业务日查看 午餐 / 晚餐 / 午晚餐合并 视图")
    add_bullet(doc, "一张表汇总：订阅会员、单次点餐、自提、请假标记")
    add_bullet(doc, "手动标记 配送状态（午餐/晚餐可分别标记）")
    add_bullet(doc, "导出 Excel 配送清单，供厨房与配送使用")
    add_bullet(doc, "一键推送顺丰：预览停靠点 → 批量推单")

    add_para(doc, "配送区域管理", bold=True)
    add_bullet(doc, "在地图上 绘制配送片区（多边形围栏）")
    add_bullet(doc, "会员地址 自动匹配 所属片区")
    add_bullet(doc, "片区地图总览：会员分布可视化")

    add_para(doc, "配送资质检验", bold=True)
    add_bullet(doc, "输入任意地址，即时判断是否在 本店服务范围 内")

    add_para(doc, "顺丰订单监控", bold=True)
    add_bullet(doc, "查看顺丰推单记录与状态")
    add_bullet(doc, "重试失败推单、取消推单、同步履约状态")
    add_bullet(doc, "导出推单报表")

    add_para(doc, "配送员管理", bold=True)
    add_bullet(doc, "增删改配送员信息")
    add_bullet(doc, "重置登录 PIN 码")
    add_bullet(doc, "查看 待结算 / 已结算 配送费")
    shots.insert(doc, "05_delivery")

    add_title(doc, "5.7 财务中心（店主专属）", 2)
    add_bullet(doc, "实收汇总：开卡工单、周卡、月卡、次卡、单次点餐、商城、退款等分项")
    add_bullet(doc, "支持 按月 / 按日 钻取明细")
    add_bullet(doc, "今日已缴开卡工单 流水列表")
    shots.insert(doc, "05_finance")

    add_title(doc, "5.8 菜单管理", 2)
    add_table(doc, ["模块", "功能"], [
        ["菜品管理", "维护菜品库：名称、图片、分类、成分、价格等"],
        ["本周菜单", "按周排期：每天午餐/晚餐安排哪道菜"],
        ["日总份库存", "设置某日某菜的可售总份数，防止超卖"],
        ["普通商品管理", "零售分类与 SKU（非套餐菜的其他商品）"],
        ["菜品分类", "分类树维护"],
    ])
    shots.insert(doc, "05_weekly_menu")

    add_title(doc, "5.9 小程序营销", 2)
    add_table(doc, ["模块", "功能"], [
        ["首页 Banner", "轮播图与跳转链接配置"],
        ["进入弹窗海报", "用户打开小程序时的活动海报"],
        ["优惠券管理", "创建券种：代金券、适用范围（购卡/单餐/商城）、有效期"],
        ["优惠券发放", "向指定会员发券、批量发放、撤销"],
        ["抖音商品设置", "抖音团购商品与本地权益（卡/券）映射"],
        ["核销记录查询", "抖音验券兑换流水"],
    ])
    shots.insert(doc, "05_marketing")

    add_title(doc, "5.10 门店配置（店主专属）", 2)
    add_bullet(doc, "门店名称、Logo、联系电话、坐标")
    add_bullet(doc, "顺丰自动推单 开关（每日定时自动推送配送大表）")
    add_bullet(doc, "配送员 计费规则：基础费 + 每多一份加价")
    add_bullet(doc, "请假截止时间等业务参数")
    add_bullet(doc, "微信支付、顺丰、抖音、地图等 第三方对接配置")

    add_title(doc, "5.11 系统管理", 2)
    add_table(doc, ["模块", "功能"], [
        ["会员卡管理", "自定义卡包模板：名称、餐次、价格、封面图、午/晚餐覆盖、有效期"],
        ["菜品分类", "全局菜品分类维护"],
        ["图片上传", "运营素材上传管理"],
    ])

    # 六、平台级能力
    add_title(doc, "六、平台级能力（多门店 / 多品牌）", 1)
    add_para(doc, "系统采用 多租户架构，一套系统可服务多个品牌、多个门店：")
    add_bullet(doc, "一个平台 下可管理 多个租户（品牌）")
    add_bullet(doc, "每个租户下可开设 多个门店")
    add_bullet(doc, "会员、菜单、订单等数据 按门店隔离")
    add_bullet(doc, "平台管理员统一管理租户、门店、管理员账号与对接配置")
    add_para(doc, "适合：已有单店经验、计划扩张的连锁品牌；或希望为多个客户部署的 SaaS 运营方。")

    # 七、支付与配送方式
    add_title(doc, "七、支付与配送方式", 1)
    add_title(doc, "7.1 支付方式", 2)
    add_table(doc, ["方式", "适用场景"], [
        ["微信支付", "小程序内购卡、单次点餐、商城下单"],
        ["会员餐次余额", "会员单次点餐直接扣次，无需再次支付"],
        ["线下 / 其他渠道", "开卡工单登记微信转账、支付宝、抖音等缴费"],
    ])

    add_title(doc, "7.2 配送履约方式", 2)
    add_table(doc, ["方式", "适用场景"], [
        ["门店自有配送员", "订阅大表、单次、商城均可指派"],
        ["顺丰同城", "大表批量推单、单次/零售派单，支持定时自动推单"],
        ["UU 跑腿", "单次点餐、卡包订单派单"],
        ["门店自提", "会员选择自提，后台核销"],
    ])

    # 八、系统自动化能力
    add_title(doc, "八、系统自动化能力", 1)
    add_para(doc, "系统内置多项 定时自动化任务，减少人工操作：")
    add_table(doc, ["自动化任务", "说明"], [
        ["每日凌晨", "清理过期请假标记；生效预约份数变更"],
        ["每日上午", "对已开启的门店 自动推送顺丰配送大表"],
        ["每日下午", "低余额扫描；次日余量不足提醒"],
        ["持续运行", "超时未支付订单自动取消；购卡超时释放锁定优惠券"],
    ])

    # 九、产品亮点总结
    add_title(doc, "九、产品亮点总结", 1)
    highlights = [
        "订阅 + 散客双模式 — 同一套菜单同时服务长期订户与偶尔点餐用户",
        "午/晚餐独立运营 — 菜单、请假、余额、配送大表均可午餐/晚餐分开管理",
        "智能配送大表 — 一张表统筹订阅、单次、自提、请假，可一键推顺丰",
        "多履约通道融合 — 自有骑手 + 顺丰 + UU，管理端统一派单与监控",
        "开卡双轨统一入账 — 线下收款与小程序在线购卡进入同一会员体系",
        "灵活会员卡模板 — 不限于周/月/次，可自定义卡包名称、餐次、餐段、有效期",
        "营销闭环 — Banner、弹窗、优惠券、抖音团购验券兑权益一体化",
        "配送范围智能匹配 — 地址自动校验是否在服务区域内",
        "微信消息触达 — 送达通知、低余额续费提醒，提升留存",
        "多租户可扩展 — 一套系统服务多品牌多门店，数据隔离",
        "厨房备餐管控 — 日库存、报损、补偿、备餐量管理",
        "会员自律工具 — 请假、份数、扣次透明，契合健康餐「自律」产品定位",
    ]
    for i, h in enumerate(highlights, 1):
        add_bullet(doc, h)

    # 十、典型业务场景
    add_title(doc, "十、典型业务场景", 1)
    scenarios = [
        ("场景一：新用户首次购卡", "用户在小程序浏览会员卡 → 选择周卡/月卡 → 使用优惠券抵扣 → 微信支付 → 设置起送日期与地址 → 次日开始配送 → 收到送达微信通知。"),
        ("场景二：老会员临时请假", "会员在小程序选择「明天请假」或指定日期区间 → 系统自动从大表剔除 → 餐次不扣减 → 请假结束后自动恢复配送。"),
        ("场景三：门店每日备餐与配送", "凌晨系统自动清理过期请假 → 上午店主打开营业概览查看今日备餐量 → 导出配送大表给厨房 → 08:50 系统自动推送顺丰（若已开启）→ 配送员在小程序按片区配送 → 确认送达后自动扣次。"),
        ("场景四：线下收款开卡", "客户微信转账购卡 → 客服在后台新建开卡工单 → 登记金额与渠道 → 点击「入账」→ 会员餐次增加、计划激活 → 与线上下单体验一致。"),
        ("场景五：抖音团购引流", "用户在抖音购买团购券 → 打开 OK饭小程序验券兑换 → 获得周卡或优惠券 → 完成转化，核销记录可在后台查询。"),
    ]
    for title, desc in scenarios:
        add_para(doc, title, bold=True)
        add_para(doc, desc, indent=True)

    # 十一、交付内容说明
    add_title(doc, "十一、交付内容说明", 1)
    add_para(doc, "采购本系统，典型交付包括：")
    add_table(doc, ["交付项", "说明"], [
        ["业务系统", "云端部署的完整后端服务"],
        ["Web 管理后台", "浏览器访问的管理端"],
        ["微信小程序", "会员端 + 配送员端（需绑定您的微信商户号与小程序）"],
        ["初始配置", "门店信息、会员卡模板、配送区域等基础配置"],
        ["对接支持", "微信支付、顺丰同城、抖音团购、高德地图等第三方对接协助"],
    ])
    add_para(doc, "具体部署方式（云端托管 / 私有化部署）、定制开发范围、售后服务条款，可在商务洽谈中进一步确认。")

    # 十二、联系我们
    add_title(doc, "十二、联系我们", 1)
    add_para(doc, "如需 产品演示、报价方案、定制需求沟通，欢迎与我们联系。")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("本文档基于 OK饭系统当前已实现功能整理，功能随版本迭代可能有所更新。正式采购以双方确认的《功能清单》与《服务协议》为准。")
    set_font(r, size=Pt(9), color=COLOR_MUTED)

    doc._shots_summary = shots.summary()  # type: ignore[attr-defined]
    return doc


def main():
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    shots = ScreenshotLoader()
    doc = build_document(shots)
    out = OUT_PATH
    try:
        doc.save(str(out))
    except PermissionError:
        out = OUT_PATH.with_stem(OUT_PATH.stem + "（含截图）")
        doc.save(str(out))
        print("原文件被占用，已另存为:")
    print(f"已生成: {out}")
    print(getattr(doc, "_shots_summary", shots.summary()))
    if shots.missing:
        print("\n提示: 缺失的截图可运行以下命令补充：")
        print("  管理后台: python scripts/capture_admin_screenshots.py")
        print("  小程序: 从微信开发者工具截图，保存到 docs/sales/screenshots/ 对应文件名")


if __name__ == "__main__":
    main()
