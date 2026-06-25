#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 OK饭 一页纸产品简介：Word 文档 + 长图 PNG。"""

from __future__ import annotations

import textwrap
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "sales"

# ── 一页纸文案 ──────────────────────────────────────────────
TITLE = "OK饭（OK Fine）健康餐配送管理系统"
SUBTITLE = "订阅配送 · 单次点餐 · 会员运营 · 智能履约  一体化解决方案"
TAGLINE = "专为健康餐 / 轻食 / 减脂餐门店打造的全链路数字化经营系统"

SECTIONS = [
    {
        "title": "产品组成",
        "lines": [
            "微信小程序（会员端 + 配送员端）— 用户点餐购卡、骑手接单配送",
            "Web 管理后台 — 店主 / 客服 / 财务日常经营与数据管理",
            "云端业务系统 — 订单、会员、配送、财务数据自动流转",
        ],
    },
    {
        "title": "核心功能一览",
        "bullets": [
            ("会员小程序", "微信一键登录 · 周/月/次卡在线购卡 · 本周菜单浏览 · 单次点餐 · 请假/改份数 · 多地址管理 · 优惠券 · 抖音团购验券 · 送达/续费消息提醒"),
            ("配送员小程序", "按片区查看当日任务 · 一键拨号/导航 · 确认送达自动扣次 · 配送费自动累计"),
            ("管理后台", "营业概览 · 会员档案 · 开卡工单 · 订单管理 · 智能配送大表 · 菜单/库存 · 财务实收 · 营销运营 · 门店配置"),
        ],
    },
    {
        "title": "六大亮点",
        "bullets": [
            ("订阅 + 散客双模式", "长期订户与偶尔点餐用户共用一套菜单与订单体系"),
            ("智能配送大表", "一张表统筹订户、单次、自提、请假，可导出 Excel、一键推顺丰"),
            ("多履约通道", "自有骑手 + 顺丰同城 + UU 跑腿 + 门店自提，统一管理"),
            ("开卡双轨入账", "小程序在线购卡与线下收款开卡工单，统一进入会员体系"),
            ("午/晚餐独立运营", "菜单、请假、余额、大表均可午餐/晚餐分开管理"),
            ("多门店可扩展", "一套系统服务多品牌、多门店，数据隔离，适合连锁扩张"),
        ],
    },
    {
        "title": "适用客户",
        "lines": [
            "健康餐 / 轻食 / 减脂餐 周订、月订门店",
            "同时服务长期订户与散客单次点餐的中央厨房",
            "计划多店扩张的连锁品牌，或 SaaS 多租户运营方",
        ],
    },
    {
        "title": "典型场景",
        "lines": [
            "新用户小程序购卡 → 设置起送日 → 每日自动配送 → 微信送达通知",
            "会员临时请假 → 大表自动剔除 → 餐次不扣减 → 到期自动恢复",
            "店主查看备餐概览 → 导出配送清单 → 顺丰自动推单 → 骑手确认送达",
            "抖音团购券 → 小程序验券兑换 → 获得会员卡/优惠券 → 完成引流转化",
        ],
    },
]

FOOTER = "如需产品演示、报价方案或定制需求，欢迎联系我们 · 功能以正式交付清单为准"


def _set_doc_font(run, name: str = "微软雅黑", size: Pt | None = None, bold: bool = False, color: RGBColor | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def generate_docx(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    _set_doc_font(r, size=Pt(18), bold=True, color=RGBColor(0x1A, 0x7A, 0x4A))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SUBTITLE)
    _set_doc_font(r, size=Pt(10), color=RGBColor(0x55, 0x55, 0x55))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TAGLINE)
    _set_doc_font(r, size=Pt(10), color=RGBColor(0x33, 0x33, 0x33))

    doc.add_paragraph()

    for sec in SECTIONS:
        p = doc.add_paragraph()
        r = p.add_run(sec["title"])
        _set_doc_font(r, size=Pt(12), bold=True, color=RGBColor(0x1A, 0x7A, 0x4A))
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)

        if "lines" in sec:
            for line in sec["lines"]:
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(line)
                _set_doc_font(r, size=Pt(9))
                p.paragraph_format.space_after = Pt(1)

        if "bullets" in sec:
            for label, desc in sec["bullets"]:
                p = doc.add_paragraph()
                r1 = p.add_run(f"{label}：")
                _set_doc_font(r1, size=Pt(9), bold=True)
                r2 = p.add_run(desc)
                _set_doc_font(r2, size=Pt(9))
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.left_indent = Cm(0.3)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(FOOTER)
    _set_doc_font(r, size=Pt(8), color=RGBColor(0x88, 0x88, 0x88))

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    print(f"Word 已生成: {path}")


def _load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
    ]
    for fp in candidates:
        if Path(fp).exists():
            try:
                return ImageFont.truetype(fp, size)
            except OSError:
                continue
    return ImageFont.load_default()


def _wrap_text(text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    lines: list[str] = []
    for paragraph in text.split("\n"):
        if not paragraph:
            lines.append("")
            continue
        current = ""
        for ch in paragraph:
            test = current + ch
            bbox = font.getbbox(test)
            if bbox[2] - bbox[0] <= max_width:
                current = test
            else:
                if current:
                    lines.append(current)
                current = ch
        if current:
            lines.append(current)
    return lines


def _draw_rounded_rect(draw: ImageDraw.ImageDraw, xy: tuple, radius: int, fill: tuple):
    draw.rounded_rectangle(xy, radius=radius, fill=fill)


def generate_long_image(path: Path, width: int = 750) -> None:
    # 配色
    BG = (0xF7, 0xFA, 0xF8)
    PRIMARY = (0x1A, 0x7A, 0x4A)
    PRIMARY_LIGHT = (0xE8, 0xF5, 0xEE)
    TEXT = (0x2C, 0x2C, 0x2C)
    TEXT_MUTED = (0x66, 0x66, 0x66)
    WHITE = (255, 255, 255)

    pad_x = 40
    content_w = width - pad_x * 2

    font_title = _load_font(36, bold=True)
    font_sub = _load_font(22)
    font_tag = _load_font(20)
    font_sec = _load_font(26, bold=True)
    font_body = _load_font(22)
    font_label = _load_font(22, bold=True)
    font_footer = _load_font(18)

    line_h = 34
    sec_gap = 28

    # 预计算高度
    height = 60  # top pad
    height += 50 + line_h + 8 + line_h + 16 + line_h + 40  # header

    for sec in SECTIONS:
        height += 36 + sec_gap // 2  # section title
        if "lines" in sec:
            for line in sec["lines"]:
                wrapped = _wrap_text(f"• {line}", font_body, content_w - 20)
                height += len(wrapped) * line_h + 6
        if "bullets" in sec:
            for label, desc in sec["bullets"]:
                text = f"{label}：{desc}"
                wrapped = _wrap_text(text, font_body, content_w - 20)
                height += len(wrapped) * line_h + 10
        height += sec_gap

    height += 60 + 30  # footer
    height += 40  # bottom pad

    img = Image.new("RGB", (width, height), BG)
    draw = ImageDraw.Draw(img)

    y = 40

    # 顶部色块
    _draw_rounded_rect(draw, (pad_x, y, width - pad_x, y + 8), 4, PRIMARY)
    y += 24

    # 标题区
    draw.text((pad_x, y), TITLE, fill=PRIMARY, font=font_title)
    y += 50
    draw.text((pad_x, y), SUBTITLE, fill=TEXT_MUTED, font=font_sub)
    y += line_h + 8
    draw.text((pad_x, y), TAGLINE, fill=TEXT, font=font_tag)
    y += line_h + 24

    # 产品组成卡片
    _draw_rounded_rect(draw, (pad_x, y, width - pad_x, y + 120), 12, PRIMARY_LIGHT)
    card_y = y + 16
    draw.text((pad_x + 16, card_y), "产品组成", fill=PRIMARY, font=font_sec)
    card_y += 36
    for line in SECTIONS[0]["lines"]:
        draw.text((pad_x + 16, card_y), f"▸ {line}", fill=TEXT, font=font_body)
        card_y += line_h
    y += 120 + sec_gap

    # 其余区块
    for sec in SECTIONS[1:]:
        draw.text((pad_x, y), sec["title"], fill=PRIMARY, font=font_sec)
        y += 36

        if "lines" in sec:
            for line in sec["lines"]:
                wrapped = _wrap_text(f"• {line}", font_body, content_w - 20)
                for wl in wrapped:
                    draw.text((pad_x + 8, y), wl, fill=TEXT, font=font_body)
                    y += line_h
                y += 6

        if "bullets" in sec:
            for label, desc in sec["bullets"]:
                # 标签加粗效果：先画标签再画正文
                prefix = f"{label}："
                prefix_w = font_label.getbbox(prefix)[2]
                wrapped = _wrap_text(desc, font_body, content_w - 20 - prefix_w)
                draw.text((pad_x + 8, y), prefix, fill=PRIMARY, font=font_label)
                if wrapped:
                    draw.text((pad_x + 8 + prefix_w, y), wrapped[0], fill=TEXT, font=font_body)
                    y += line_h
                    for wl in wrapped[1:]:
                        draw.text((pad_x + 8, y), wl, fill=TEXT, font=font_body)
                        y += line_h
                y += 10

        y += sec_gap

    # 底部分隔线
    draw.line([(pad_x, y), (width - pad_x, y)], fill=(0xCC, 0xDD, 0xD4), width=2)
    y += 20

    footer_lines = _wrap_text(FOOTER, font_footer, content_w)
    for fl in footer_lines:
        bbox = font_footer.getbbox(fl)
        tw = bbox[2] - bbox[0]
        draw.text(((width - tw) // 2, y), fl, fill=TEXT_MUTED, font=font_footer)
        y += 28

    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(str(path), format="PNG", optimize=True)
    print(f"长图已生成: {path}  ({width}x{height}px)")


def main():
    docx_path = OUT_DIR / "OK饭一页纸产品简介.docx"
    png_path = OUT_DIR / "OK饭一页纸产品简介.png"
    generate_docx(docx_path)
    generate_long_image(png_path)
    print("\n输出目录:", OUT_DIR.resolve())


if __name__ == "__main__":
    main()
